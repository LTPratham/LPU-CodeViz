import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color="0A2540", size="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color)
    tcBorders.append(left)
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'nil')
        tcBorders.append(node)
    tcPr.append(tcBorders)

def style_table(table, header_bg="0A2540", row_shd="F5F7FA"):
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
    for r_idx, row in enumerate(table.rows[1:], start=1):
        shd_color = row_shd if r_idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            if shd_color != "FFFFFF":
                set_cell_background(cell, shd_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), 'E2E8F0')
            tcBorders.append(bottom)
            for b in ['left', 'right', 'top']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'nil')
                tcBorders.append(node)
            tcPr.append(tcBorders)

def add_custom_heading(doc, text, level, space_before=18, space_after=6):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    run = h.add_run(text)
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(10, 37, 64)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 128, 128)
        run.bold = True
    return h

def add_body_paragraph(doc, text, bold_prefix="", space_after=8, list_style=None):
    p = doc.add_paragraph(style=list_style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
        r_prefix.font.name = 'Calibri'
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = RGBColor(10, 37, 64)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_callout_box(doc, text, title="FINANCIAL NOTICE", color="D97706"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "FDFDF9")
    set_cell_left_border(cell, color=color, size="36")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"{title}\n")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(217, 119, 6)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = RGBColor(51, 51, 51)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(12)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def generate_financials():
    doc = Document()
    
    # Page Geometry Setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Normal style override
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # ─── COVER PAGE ───
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("MARKET SIZING (TAM/SAM/SOM)\n& FINANCIAL PROJECTIONS")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.color.rgb = RGBColor(10, 37, 64)
    run_title.bold = True
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(80)
    run_sub = sub_p.add_run("Commercial Feasibility, 5-Year Scaling, Operating Cost Analysis & Unit Economics")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0, 128, 128)
    run_sub.bold = True
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(20)
    run_meta = meta_p.add_run("Author: Prathamesh Sawarkar (Registration ID: 12509401)\n"
                              "Scope: CodeCanvas (LPU CodeViz)\n"
                              "Institutional Focus: B2B Higher-Education SaaS\n"
                              "Last Updated: August 11, 2026")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(100, 110, 120)
    
    doc.add_page_break()
    
    # Setup Header & Footer for rest of document
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header_p.add_run("CodeCanvas — Market & Financial Modeling")
    run_h.font.name = 'Arial'
    run_h.font.size = Pt(8.5)
    run_h.font.color.rgb = RGBColor(120, 130, 140)
    
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run("Page ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9.0)
    run_f.font.color.rgb = RGBColor(120, 130, 140)
    add_page_number(run_f)
    
    # ─── SECTION 1: PRICING STRUCTURE ───
    add_custom_heading(doc, "1. Commercial Pricing Structure (SAC Code 997331)", level=1)
    
    add_body_paragraph(doc, 
        "CodeCanvas utilizes a hybrid commercial model, combining B2C individual upgrades with B2B departmental and institutional site licenses. All prices are structured to comply with Indian GST rules under SAC Code 997331 (Licensing services for the right to use computer software).",
        bold_prefix="Pricing Rationale: "
    )
    
    table_price = doc.add_table(rows=4, cols=5)
    table_price.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_price.autofit = False
    table_price.columns[0].width = Inches(1.3)
    table_price.columns[1].width = Inches(1.1)
    table_price.columns[2].width = Inches(1.0)
    table_price.columns[3].width = Inches(1.1)
    table_price.columns[4].width = Inches(2.0)
    
    hdr_pri = table_price.rows[0].cells
    hdr_pri[0].text = "License Tier"
    hdr_pri[1].text = "Base Price"
    hdr_pri[2].text = "GST (18%)"
    hdr_pri[3].text = "Total Price"
    hdr_pri[4].text = "SaaS Capabilities Included"
    
    price_rows = [
        ("Individual Student Monthly", "Rs. 254.24", "Rs. 45.76", "Rs. 300.00", "Visualizer sandbox, AI Tutor Chatbot, personal trace history Dashboard, and streaks."),
        ("Standard Departmental Annual", "Rs. 84,745.76", "Rs. 15,254.24", "Rs. 1,00,000.00", "500 student seats, 10 teacher licenses, classroom boards, auto-grader, and NAAC reports."),
        ("Institutional Enterprise Annual", "Rs. 2,54,237.28", "Rs. 45,762.72", "Rs. 3,00,000.00", "Unlimited seats, LMS integration, advanced telemetry, SSO, backups, support manager.")
    ]
    for idx, (tier, base, gst, tot, cap) in enumerate(price_rows, start=1):
        row = table_price.rows[idx].cells
        row[0].text = tier
        row[1].text = base
        row[2].text = gst
        row[3].text = tot
        row[4].text = cap
    style_table(table_price)
    
    doc.add_page_break()
    
    # ─── SECTION 2: TAM/SAM/SOM ───
    add_custom_heading(doc, "2. Market Sizing: TAM / SAM / SOM", level=1)
    
    add_body_paragraph(doc, 
        "Our market opportunity spans technical education institutions globally, narrowing down to the Indian engineering ecosystem as our initial launching target:",
        bold_prefix="Market Definition: "
    )
    
    add_body_paragraph(doc, 
        "Total Addressable Market (TAM): Globally, there are approximately 50,000 higher education institutions teaching computer science, software development, or database systems. Assuming an average annual institutional site license contract value of $4,000 USD (approx. Rs. 3,30,000), our Total Addressable Market stands at $200 Million USD (approx. Rs. 1,650 Crores).",
        bold_prefix="• ", list_style="Normal"
    )
    
    add_body_paragraph(doc, 
        "Serviceable Addressable Market (SAM): In India, there are 8,500 institutions offering technical higher education (B.Tech CS/IT, BCA, and MCA degrees), including 3,500 AICTE-approved engineering colleges and 5,000 other accredited universities. With a target average institutional contract value of Rs. 1,50,000 per college annually, our Serviceable Addressable Market is Rs. 127.5 Crores ($15.3 Million USD).",
        bold_prefix="• ", list_style="Normal"
    )
    
    add_body_paragraph(doc, 
        "Serviceable Obtainable Market (SOM): Within 5 years, we aim to capture B2B contracts with 350 Tier-1 and Tier-2 engineering colleges and universities in India (focusing initially on Punjab, Haryana, Karnataka, Tamil Nadu, and Maharashtra). Assuming a B2B Average Contract Value (ACV) of Rs. 2,00,000, our Serviceable Obtainable Market is Rs. 7.0 Crores ($840,000 USD) ARR.",
        bold_prefix="• ", list_style="Normal"
    )
    
    add_callout_box(doc, 
        "TAM (Global CS Higher Ed): Rs. 1,650 Crores ($200M USD)\n"
        "SAM (Indian CS Higher Ed): Rs. 127.5 Crores ($15.3M USD)\n"
        "SOM (Indian Target Colleges): Rs. 7.0 Crores ($840K USD) ARR",
        title="MARKET MATRIX SIZING SUMMARY", color="0A2540"
    )
    
    # ─── SECTION 3: 5-YEAR PROJECTIONS ───
    add_custom_heading(doc, "3. 5-Year Revenue Projections Matrix", level=1)
    
    table_proj = doc.add_table(rows=6, cols=6)
    table_proj.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_proj.autofit = False
    table_proj.columns[0].width = Inches(1.8)
    table_proj.columns[1].width = Inches(0.8)
    table_proj.columns[2].width = Inches(0.8)
    table_proj.columns[3].width = Inches(0.9)
    table_proj.columns[4].width = Inches(1.0)
    table_proj.columns[5].width = Inches(1.2)
    
    hdr_pro = table_proj.rows[0].cells
    hdr_pro[0].text = "Metric Tracker"
    hdr_pro[1].text = "Year 1"
    hdr_pro[2].text = "Year 2"
    hdr_pro[3].text = "Year 3"
    hdr_pro[4].text = "Year 4"
    hdr_pro[5].text = "Year 5"
    
    projs = [
        ("B2B College Signups", "5", "25", "75", "180", "350"),
        ("B2C Student Upgrades", "1,000", "5,000", "15,000", "30,000", "60,000"),
        ("B2B Revenue (INR)", "Rs. 10 Lakhs", "Rs. 50 Lakhs", "Rs. 1.5 Crores", "Rs. 3.66 Crores", "Rs. 7.0 Crores"),
        ("B2C Revenue (INR)", "Rs. 3 Lakhs", "Rs. 15 Lakhs", "Rs. 45 Lakhs", "Rs. 90 Lakhs", "Rs. 1.8 Crores"),
        ("Total ARR (INR)", "Rs. 13 Lakhs", "Rs. 65 Lakhs", "Rs. 1.95 Crores", "Rs. 4.56 Crores", "Rs. 8.8 Crores")
    ]
    for idx, row_data in enumerate(projs, start=1):
        row = table_proj.rows[idx].cells
        for col_idx, text in enumerate(row_data):
            row[col_idx].text = text
    style_table(table_proj)
    
    doc.add_page_break()
    
    # ─── SECTION 4: OPEX MODEL ───
    add_custom_heading(doc, "4. Operating Cost Model (OpEx) & Unit Economics", level=1)
    
    add_body_paragraph(doc, 
        "CodeCanvas operates as a highly scalable cloud product. By delegating program evaluation logic to a fast, pay-per-token API (Groq) rather than hosting expensive GPU-enabled virtual servers, we minimize operating overhead.",
        bold_prefix="Infrastructure Efficiency: "
    )
    
    add_body_paragraph(doc, 
        "1. Cloud Infrastructure: Next.js hosting via Vercel costs Rs. 24,000 in Year 1, rising to Rs. 4,80,000 in Year 5. Supabase enterprise logs and database sizing starts at Rs. 30,000 in Year 1 and expands to Rs. 18,00,000 in Year 5 as connection volumes scale.",
        list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "2. Groq AI Token Calculations: Each active student is estimated to execute an average of 40 visual traces per month in lab classes. One visual trace averages 10,000 tokens (8K prompt + 2K output). Under Groq rates ($0.59/M input, $0.79/M output), one trace costs approx. Rs. 0.50 INR. This equates to Rs. 20.00/month per active student.",
        list_style="List Bullet"
    )
    
    add_callout_box(doc, 
        "The marginal infrastructure and API cost per active student (Rs. 20) represents only 7.8% of the B2C Student monthly base license fee (Rs. 254.24), yielding an outstanding Gross Margin of 92.2%.",
        title="UNIT ECONOMICS VERIFICATION", color="008080"
    )
    
    # ─── SECTION 5: B2B UNIT ECONOMICS ───
    add_custom_heading(doc, "5. B2B Unit Economics & Sales CAC/LTV", level=1)
    
    add_body_paragraph(doc, 
        "For B2B institutional sales, we track the following unit economics to justify marketing expenditure:"
    )
    
    add_body_paragraph(doc, 
        "Customer Acquisition Cost (CAC): Rs. 30,000 per college (covering on-campus marketing sales rep visits, catalogs, and targeted institutional social media outreach).",
        bold_prefix="• ", list_style="Normal"
    )
    add_body_paragraph(doc, 
        "Average Contract Value (ACV): Rs. 2,00,000 per year per college.",
        bold_prefix="• ", list_style="Normal"
    )
    add_body_paragraph(doc, 
        "Customer Lifetime Value (LTV): Rs. 6,00,000 per college (calculated over an average retention period of 3 years).",
        bold_prefix="• ", list_style="Normal"
    )
    add_body_paragraph(doc, 
        "LTV / CAC Ratio: 20.0x (An LTV/CAC ratio of >3.0x is considered highly viable for institutional SaaS products, representing extremely efficient sales economics).",
        bold_prefix="• ", list_style="Normal"
    )
    
    # ─── SECTION 6: BREAK-EVEN ───
    add_custom_heading(doc, "6. Break-Even Analysis", level=1)
    
    add_body_paragraph(doc, 
        "Fixed operating costs in Year 1 are estimated at Rs. 80,000 per month (includes domain registration, core hosting tiers, travel budgets for sales meetings, and advertising). With an average gross margin of 85% on B2B licenses, the monthly revenue threshold required to break even is calculated as follows:"
    )
    add_body_paragraph(doc, 
        "Break-Even Revenue = Rs. 80,000 / 0.85 = Rs. 94,117 / month.",
        bold_prefix="Formula: "
    )
    add_body_paragraph(doc, 
        "To satisfy this threshold, CodeCanvas needs to maintain either: (a) 6 active B2B Departmental Annual subscriptions (Rs. 1,00,000/year each), or (b) 2 active B2B Institutional Enterprise subscriptions (Rs. 3,00,000/year each). Break-even is projected to occur in Month 4 of Year 1 operations."
    )
    
    # Save the document
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "market_and_financials.docx")
    doc.save(out_path)
    print(f"Financial projections document generated successfully at: {out_path}")

if __name__ == "__main__":
    generate_financials()
