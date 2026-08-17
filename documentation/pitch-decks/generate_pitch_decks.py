import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color="0A2540", size="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color)
    tcBorders.append(left)
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'nil')
        tcBorders.append(node)
    tcPr.append(tcBorders)

def style_table(table, header_bg="0A2540", row_shd="F5F7FA"):
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
    for r_idx, row in enumerate(table.rows[1:], start=1):
        shd_color = row_shd if r_idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            if shd_color != "FFFFFF":
                set_cell_background(cell, shd_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), 'E2E8F0')
            tcBorders.append(bottom)
            for b in ['left', 'right', 'top']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'nil')
                tcBorders.append(node)
            tcPr.append(tcBorders)

def add_custom_heading(doc, text, level, space_before=18, space_after=6):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    run = h.add_run(text)
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(10, 37, 64)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 128, 128)
        run.bold = True
    return h

def add_body_paragraph(doc, text, bold_prefix="", space_after=8, list_style=None):
    p = doc.add_paragraph(style=list_style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
        r_prefix.font.name = 'Calibri'
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = RGBColor(10, 37, 64)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_callout_box(doc, text, title="VISUAL SUGGESTION", color="008080"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4FBFB")
    set_cell_left_border(cell, color=color, size="36")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"{title}\n")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(9.5)
    run_title.font.color.rgb = RGBColor(0, 128, 128)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(51, 51, 51)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(10)

def add_speaker_notes(doc, notes_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_left_border(cell, color="64748B", size="24")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    
    r_title = p.add_run("SPEAKER TALKING POINTS\n")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(9.0)
    r_title.font.color.rgb = RGBColor(100, 116, 139)
    
    r_notes = p.add_run(notes_text)
    r_notes.font.name = 'Calibri'
    r_notes.font.size = Pt(10.0)
    r_notes.font.color.rgb = RGBColor(71, 85, 105)
    r_notes.italic = True
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(14)

def setup_page_geometry(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)

def add_header_footer(doc, title_text):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header_p.add_run(title_text)
    run_h.font.name = 'Arial'
    run_h.font.size = Pt(8.5)
    run_h.font.color.rgb = RGBColor(120, 130, 140)
    
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run("Page ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9.0)
    run_f.font.color.rgb = RGBColor(120, 130, 140)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run_f._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def generate_technical_deck():
    doc = Document()
    setup_page_geometry(doc)
    
    # Title Cover Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("TECHNICAL PITCH DECK\nFOR JUDGES")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.color.rgb = RGBColor(10, 37, 64)
    run_title.bold = True
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(80)
    run_sub = sub_p.add_run("CodeCanvas: Dynamic Code Visualization & Simulated Execution Engine")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0, 128, 128)
    run_sub.bold = True
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(20)
    run_meta = meta_p.add_run("Author: Prathamesh Sawarkar (Registration ID: 12509401)\n"
                              "Scope: LPU Computer Science Department\n"
                              "Stack: Next.js + FastAPI + Supabase PostgreSQL + Groq AI Tracing\n"
                              "Contact: prathameshsawarkar1@gmail.com")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(100, 110, 120)
    
    doc.add_page_break()
    add_header_footer(doc, "CodeCanvas — Technical Pitch Deck (v2.0)")
    
    slides = [
        ("Slide 1: Title & Technical Overview",
         "Showcase a dark-themed visual block tracing code inputs to a graphical data structures map.",
         "• Title: CodeCanvas: Visualizing Code State\n"
         "• Subtitle: An AI-Simulated Code Tracing and Interactive Tutoring Architecture\n"
         "• Stack: Next.js App Router | FastAPI | Supabase PostgreSQL | Groq Llama-3.3\n"
         "• Scope: Lovely Professional University Computer Science Department",
         "Hello, judges. Today, I am presenting the technical architecture of CodeCanvas—an interactive code tracing platform designed to solve the visualization gap in computer science labs. Instead of running expensive, risky sandboxes to execute student code, CodeCanvas uses a dual-layer client-server architecture powered by high-throughput AI inference models to trace execution state dynamically in under two seconds."),
        
        ("Slide 2: Technical Architecture",
         "System flowchart: Monaco Editor UI -> FastAPI (Railway) -> Groq API (Llama 3.3 70B) -> Steps JSON -> Visual Canvas.",
         "• Client-Server Division: Frontend handles visual layers; backend manages AI-simulated execution compilation.\n"
         "• Serverless Middlewares: FastAPI functions as API broker; Supabase handles authentications & SQL variables.\n"
         "• Speed Goals: Render loops compile client-side using Framer Motion to bypass browser layout bottlenecks.",
         "The standard way to build a code execution platform like LeetCode is to run student code in isolated, virtualized Docker containers on your servers. However, this is expensive and poses serious security risks. CodeCanvas bypasses this by introducing AI-Simulated Tracing. Let's look at the flow: the student writes code, our FastAPI server packages the file, Groq evaluates it, and the frontend renders the state transition snapshots."),
        
        ("Slide 3: The AI-Simulated Tracing Engine",
         "Illustration showing source code passing through a filtration layer to a virtual execution AI model, yielding strict JSON arrays.",
         "• Core Model: Llama-3.3-70b-versatile hosted on Groq for sub-2.0s generation times.\n"
         "• Prompt Construction: System directives force strict JSON snapshots detailing lines, scopes, variables, and explanations.\n"
         "• Safety Parameters: Interprets infinite loops and exceptions safely without execution halts.",
         "Our tracing backend does not execute the student's code on a physical processor. Instead, the code is sent to the Llama-3.3 model via Groq. The model acts as a virtual execution engine, outputting a strict JSON array representing every execution step. This is fast, secure, and generates natural-language explanations for every variable mutation."),
        
        ("Slide 4: Frontend State Rendering & Monaco Integration",
         "Mockup of Monaco Editor with highlighted code line next to a visual canvas with elements (e.g. array bars sorting).",
         "• Code Input: Embeds Monaco Editor for autocomplete, formatting, and standard theme overrides.\n"
         "• Animations: Framer Motion hooks transition elements smoothly at 60 fps.\n"
         "• Player Interface: Video-style playback controls (Next, Prev, Play, Pause, Speed slider) drive the state indexes.",
         "On the client side, we embed the Monaco Editor, giving students a professional development environment. The visualizer dashboard reads the trace JSON and passes it to Framer Motion components. As the student clicks 'Next', the variables and canvas states transition smoothly at 60 frames per second, matching the exact line highlight in the editor."),
        
        ("Slide 5: Database Models & Row-Level Security (RLS)",
         "Relational schema chart linking user profiles to classrooms, enrollments, assignments, and submissions.",
         "• Engine: Supabase PostgreSQL with automated trigger bindings.\n"
         "• Row-Level Security (RLS): Strict policies isolate student data; teachers query rosters using invite keys.\n"
         "• Logging Index: Logs telemetry of trace history and streaks, driving institutional dashboards.",
         "Our database is hosted on Supabase PostgreSQL. To protect student data and support B2B multi-tenancy, we enabled Row-Level Security on all tables. Students can only query their own submissions and profiles, while professors can view enrollments and progress logs linked to their class invite codes."),
        
        ("Slide 6: Core API Contract (POST /api/trace)",
         "Side-by-side payload view showing request parameters and step variables JSON output.",
         "• Request: code (string), lang (string), data_structure (string).\n"
         "• State Snapshots: Contains line_number (integer), explanation (string), variables (key-value object), and data_structure_state (custom visualization highlights).\n"
         "• Schema Compliance: Fails gracefully if output lacks required snapshot parameters.",
         "This is our primary API contract. When calling /api/trace, the payload contains the code string and target topic. The engine returns a step snapshot with the line number, natural-language explanation, variable states, and data structure mutations—such as array states and indices to swap—allowing the UI to animate transitions."),
        
        ("Slide 7: Security Posture: RCE Defense",
         "Diagram contrasting Docker server execution vulnerabilities (CPU attacks, kernel breaks) with safe AI-Simulated tracing.",
         "• Server Hazard: Standard IDEs compile student scripts natively, exposing servers to remote command runs.\n"
         "• AI Shielding: Malicious loops, system access attempts, and memory leaks are parsed and reported conceptually by the model, never hitting server hardware.\n"
         "• Zero Footprint: Minimizes Docker micro-virtualization maintenance costs to zero.",
         "Because we don't compile student code natively, we are secure against Remote Code Execution (RCE) attacks. If a student submits an infinite loop, system command execution, or a fork-bomb script, the AI model detects and describes the behavior in the JSON trace rather than compiling it on the server hardware. This saves system resources and eliminates sandbox vulnerabilities."),
        
        ("Slide 8: Technical Performance & Latency SLAs",
         "Latency stats graph showing average response speed distributions and cache hit benefits.",
         "• Response Speed: Trace generation resolves in <= 2.0s at the 95th percentile.\n"
         "• Caching Layer: Database caching of pre-loaded syllabus templates yields >= 45% hit ratio with sub-100ms response.\n"
         "• Session Scale: Serverless hosting scales automatically to support >= 5,000 concurrent student connections.",
         "Performance is key in lab environments. Our AI tracing calls average under 1.5 seconds. For standard templates pre-loaded into our syllabus module, we use a database caching layer. When a student visualizes a template example, the system retrieves it from cache instantly, delivering a sub-100ms response."),
        
        ("Slide 9: Gamification & Certification Architecture",
         "Workflow chart: Daily activity -> Streak validation -> XP triggers -> Cryptographic verify link.",
         "• Streak Engine: Postgres SQL trigger functions recalculate streak values on active days.\n"
         "• Certification: Automated PDF generator compiles signed certifications with verify slugs (/verify/[id]).\n"
         "• Outcomes Evidence: Aggregated trace telemetry forms proof of lab participation for audits.",
         "To keep students motivated, we integrated a database-driven gamification engine. Tracing code triggers database RPC functions to update XP and streaks. Once a topic is complete, the platform generates a cryptographically signed PDF certificate. The certificate carries a unique ID verify route, providing proof of outcome attainment for university files."),
        
        ("Slide 10: Future Roadmap & Integrations",
         "Timeline mapping LTI 1.3 LMS sync -> Pyodide WASM client offline execution -> ABC Portal link.",
         "• LTI 1.3 Compliance: Roster syncing and gradebook passback to Canvas, Blackboard, and Moodle.\n"
         "• WASM Client Engine: Integrating Pyodide directly in-browser to trace basic scripts offline.\n"
         "• ABC Integration: Sync student progress with the national Academic Bank of Credits (ABC) portal.",
         "Looking ahead, we are expanding our integrations. We are implementing LTI 1.3 standards to link rosters and gradebooks with Canvas and Moodle. We are also integrating Pyodide for WebAssembly client-side tracing and syncing progress files directly with the national Academic Bank of Credits (ABC) portal.")
    ]
    
    for title, visual, copy, notes in slides:
        add_custom_heading(doc, title, level=1)
        add_callout_box(doc, visual, title="SLIDE VISUAL LAYOUT SUGGESTION")
        add_custom_heading(doc, "Slide Text Content:", level=2)
        for line in copy.split('\n'):
            add_body_paragraph(doc, line, list_style="List Bullet")
        add_speaker_notes(doc, notes)
        doc.add_page_break()
        
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "pitch_deck_technical.docx")
    doc.save(out_path)
    print(f"Technical Pitch Deck generated successfully at: {out_path}")

def generate_problem_deck():
    doc = Document()
    setup_page_geometry(doc)
    
    # Title Cover Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("PROBLEM-ORIENTED PITCH DECK\nFOR JUDGES")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.color.rgb = RGBColor(10, 37, 64)
    run_title.bold = True
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(80)
    run_sub = sub_p.add_run("CodeCanvas: Visual Coding Education, Faculty Automation & Accreditation Compliance")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0, 128, 128)
    run_sub.bold = True
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(20)
    run_meta = meta_p.add_run("Author: Prathamesh Sawarkar (Registration ID: 12509401)\n"
                              "Scope: LPU Computer Science Department\n"
                              "Target: Higher Education B2B SaaS\n"
                              "Contact: prathameshsawarkar1@gmail.com")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(100, 110, 120)
    
    doc.add_page_break()
    add_header_footer(doc, "CodeCanvas — Problem-Oriented Pitch Deck (v2.0)")
    
    slides = [
        ("Slide 1: Title & Presentation Hook",
         "A clean showcase screenshot of the Monaco visual sandbox, highlighted in both Light and Dark modes.",
         "• Title: CodeCanvas: Visual Learning for Next-Gen Coders\n"
         "• Subtitle: Bridging the Educational Coding Gap in Higher Education\n"
         "• Key Metrics: Improving Lab Pass Rates | Automating Accreditation | Interactive AI Tutoring\n"
         "• Developer: Prathamesh Sawarkar (Registration ID: 12509401)",
         "Hello, judges. Today, I'm introducing CodeCanvas—a platform designed to make computer science labs interactive, intuitive, and compliant. We solve a major problem: students struggle to understand abstract coding concepts, while professors spend too much time on grading and accreditation paperwork. CodeCanvas brings memory models to life."),
        
        ("Slide 2: The Visualization Gap in Coding Labs",
         "Illustration contrasting blackboard lecture slides (static) vs. terminal GDB printouts (confusing for beginners).",
         "• Concepts are Invisible: Variables, pointers, recursion call stacks are invisible in standard code execution.\n"
         "• Instructors are Stretched: 1 instructor per 60+ students makes 1-on-1 execution walkthroughs impossible.\n"
         "• Debugging Friction: Beginners write syntactically correct code but lack the visual feedback to resolve logical bugs.",
         "When students start coding, they struggle to visualize what happens inside memory. Pointers, linked lists, recursive stacks, and database joins are explained on static blackboards or slides. When students write code in labs, they can't visualize variables changing, leading to logical errors. With large lab cohorts, instructors can't offer 1-on-1 tracing guidance."),
        
        ("Slide 3: Student Story — Friction Points",
         "Visual storyboard representing Aarav, a 1st year B.Tech student stuck coding at midnight.",
         "• Student Profile: Aarav, B.Tech CSE student at LPU.\n"
         "• The Problem: Stuck late at night with a recursion stack overflow while preparing for a lab exam.\n"
         "• The Consequence: No instructor available, standard error messages are confusing, leading to student dropouts.",
         "Let's put a face to this problem. Meet Aarav, a first-year student studying computer science. It's late at night, and he's coding a recursive binary search algorithm. He gets stuck in an infinite recursion loop. He has no teacher to help him, and debugging terminals are confusing. Without immediate, visual guidance, Aarav is likely to give up."),
        
        ("Slide 4: The Faculty & Dean's Administrative Burden",
         "Graphic of a lab instructor buried under printout files, roster spreadsheets, and NAAC/NBA compliance papers.",
         "• Grading Backlog: Hand-grading trace outputs for 240+ students takes up valuable faculty tutoring time.\n"
         "• Auditing Demands: Regulatory reviews require physical evidence dossiers proving outcome tracking.\n"
         "• Learner Classification: NAAC Criterion 2.2 requires tracking and targeting slow vs. fast learners.",
         "This problem affects institutions too. Engineering departments must collect extensive student data to pass audits by NAAC and NBA. Professors spend hours grading identical lab submissions and compiling physical portfolios. Identifying slow vs. fast learners remains manual and subjective."),
        
        ("Slide 5: CodeCanvas Solution",
         "High-impact graphic showing CodeCanvas features: Monaco Sandbox, 10 animated canvases, and HOD dashboards.",
         "• Interactive Memory Canvas: Visualizes variables, nodes, sorting bars, and tables dynamically in-browser.\n"
         "• 24/7 AI Coding Companion: Provides hints and debugging pointers contextually without giving solutions.\n"
         "• Automated Accreditation: One-click extraction of student telemetry mapped to criteria lists.",
         "CodeCanvas solves this by making memory visible and tutoring immediate. Our visual sandbox animates arrays, sorting, stacks, and database tables in real-time. A student can step through code line-by-line. If they get stuck, the built-in AI Tutor sidebar answers questions contextually, acting as an empathetic coding companion."),
        
        ("Slide 6: The Student User Journey",
         "Horizontal timeline: Write Code -> Animate Traces -> Inspect Variable Mutates -> Engage AI Chat -> Earn XP & Badge.",
         "• Input: Code typed in Monaco Editor (C, Python, SQL).\n"
         "• Playback: 1-click execution traces visual progression steps.\n"
         "• Hinting: AI Tutor sidebar checks active line scope, guiding the student through compile or logical errors.",
         "Here is the student's journey: Aarav types his code in our Monaco editor. Instead of running it blindly, he hits 'Visualize' to trace the states step-by-step. He sees variables mutating in the sidebar. When he is confused, he chats with the AI Tutor, resolves the bug himself, earns XP, and maintains his coding streak."),
        
        ("Slide 7: The Faculty Hub: Class Boards",
         "Mockup of the Teacher Dashboard showing enrollment rosters, assignment progress rings, and override grading panels.",
         "• Quick Setup: Deploy classrooms with unique invite codes; students enroll instantly.\n"
         "• Anti-Cheat Grader: Auto-flags templates matching unmodified files, routing validated scores to teacher panel.\n"
         "• Gradebook Overrides: Allows teachers to review step-by-step student playback runs and override grades.",
         "For faculty, we built the Multi-Tenant Classroom Board. A professor creates a class, generates an invite code, and shares it with students. They publish custom assignments aligned with their syllabus. CodeCanvas auto-grades submissions, logs student streaks, and allows professors to override scores."),
        
        ("Slide 8: Accreditation Mappings (NAAC & NBA Compliance)",
         "Infographic showing maps: Trace counts to NAAC 2.3.1 (ICT Tools) | Step metrics to NAAC 2.2.1 (Learner profiles).",
         "• NAAC Criterion 2.3.1 (ICT experiential): CodeCanvas visualizer logs prove active, student-centered learning adoption.\n"
         "• NAAC Criterion 2.2.1 (Slow vs. Fast): Identifies and tracks slow learners automatically based on lab tracing telemetry.\n"
         "• NBA Criteria (CO-PO): Connects assignments to specific Program Outcomes for certified folder exports.",
         "CodeCanvas is an administrative asset. To pass NAAC evaluations, departments can export data telemetry logs. Visualizer usage maps to NAAC Criterion 2.3 for experiential learning, while tracking trace steps identifies slow vs. fast learners for Criterion 2.2. We also map assignments to NBA Course Outcomes."),
        
        ("Slide 9: Commercial Model & Market Opportunity",
         "Pricing tables for Indian colleges, including base fees, 18% GST splits (SAC 997331), and Razorpay links.",
         "• TAM: 3,500+ engineering institutes and 1.5 million students entering Indian engineering streams annually.\n"
         "• Student Plan: Rs. 300 / month (GST incl.) - Visual sandbox, tutor, history.\n"
         "• Department Plan: Rs. 1,00,000 / year (GST incl.) - Up to 500 seats, classrooms, grading dashboards.\n"
         "• Enterprise Plan: Rs. 3,00,000 / year (GST incl.) - Unlimited seats, SSO, exports, backups.",
         "CodeCanvas operates under a SaaS B2B commercial framework. The Indian market features over 3,500 engineering colleges, representing a large addressable market. We sell individual student monthly plans, departmental annual plans (up to 500 seats), and institutional enterprise licenses. All transactions include 18% GST and route through Razorpay."),
        
        ("Slide 10: Future Roadmap",
         "Syllabus timeline: B2B Launch -> Vernacular selector modes -> LMS grade passbacks -> ABC Portal syncing.",
         "• Completed: Sandboxes, database models, class invites, trace history, and NAAC telemetry downloads.\n"
         "• In Progress: Razorpay integrations, DPDP privacy settings, and ROI calculators.\n"
         "• Horizon: Vernacular translations (Hindi/Tamil), LMS grade passback (Canvas/Moodle), and ABC Portal integration.",
         "In terms of our roadmap: we have completed the core visualizer sandbox, database migration setup, classroom portals, and NAAC export analytics. We are currently implementing Razorpay webhook integrations and privacy options. Next, we will add regional translations and sync gradebooks with LMS platforms like Canvas. Thank you, and I look after your questions.")
    ]
    
    for title, visual, copy, notes in slides:
        add_custom_heading(doc, title, level=1)
        add_callout_box(doc, visual, title="SLIDE VISUAL LAYOUT SUGGESTION")
        add_custom_heading(doc, "Slide Text Content:", level=2)
        for line in copy.split('\n'):
            add_body_paragraph(doc, line, list_style="List Bullet")
        add_speaker_notes(doc, notes)
        doc.add_page_break()
        
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "pitch_deck_problem_oriented.docx")
    doc.save(out_path)
    print(f"Problem-Oriented Pitch Deck generated successfully at: {out_path}")

if __name__ == "__main__":
    generate_technical_deck()
    generate_problem_deck()
