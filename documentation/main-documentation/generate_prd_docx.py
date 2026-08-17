import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color="0A2540", size="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Left border: thick
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size) # 36 dxa = 4.5 pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color)
    tcBorders.append(left)
    
    # Clear top, bottom, right borders
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'nil')
        tcBorders.append(node)
        
    tcPr.append(tcBorders)

def style_table(table, header_bg="0A2540", row_shd="F5F7FA"):
    # Header styling
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
    # Data rows styling
    for r_idx, row in enumerate(table.rows[1:], start=1):
        shd_color = row_shd if r_idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            if shd_color != "FFFFFF":
                set_cell_background(cell, shd_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            # Apply light horizontal border at the bottom of the cell
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4') # 0.5 pt
            bottom.set(qn('w:color'), 'E2E8F0')
            tcBorders.append(bottom)
            
            # Nil for vertical borders and top border
            for b in ['left', 'right', 'top']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'nil')
                tcBorders.append(node)
                
            tcPr.append(tcBorders)

def add_callout_box(doc, text, title="IMPORTANT NOTICE", color="D97706"):
    # Callout is represented by a 1x1 table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "FBFBFB")
    set_cell_left_border(cell, color=color, size="36")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"{title}\n")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(217, 119, 6) if color == "D97706" else RGBColor(10, 37, 64)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = RGBColor(51, 51, 51)
    
    # Add an empty paragraph after table for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(12)

def add_custom_heading(doc, text, level, space_before=18, space_after=6):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    run = h.add_run(text)
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(10, 37, 64) # Primary Accent
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 128, 128) # Secondary Accent (Teal)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(51, 51, 51) # Charcoal
        run.bold = True
        run.font.italic = True
    return h

def add_body_paragraph(doc, text, bold_prefix="", space_after=8, list_style=None):
    p = doc.add_paragraph(style=list_style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
        r_prefix.font.name = 'Calibri'
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = RGBColor(51, 51, 51)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_code_block(doc, code_text):
    # A single-cell table styled like a code block
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    # Nil borders, thin light gray left border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12') # 1.5pt
    left.set(qn('w:color'), '94A3B8')
    tcBorders.append(left)
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'nil')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.0)
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    # Spacing paragraph after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(8)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def generate_prd():
    doc = Document()
    
    # Page Geometry Setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Normal style override
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # ─── TITLE PAGE (Modern Academic Cover) ───
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(28)
    run_title.font.color.rgb = RGBColor(10, 37, 64)
    run_title.bold = True
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(48)
    run_sub = subtitle_p.add_run("CodeCanvas: Next-Generation AI-Simulated Code Visualization\n& Interactive Computer Science Tutoring Platform")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0, 128, 128)
    run_sub.bold = True
    
    badge_p = doc.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_p.paragraph_format.space_before = Pt(0)
    badge_p.paragraph_format.space_after = Pt(120)
    run_badge = badge_p.add_run("Targeted for Higher Education Engineering Colleges & Accreditation Compliance")
    run_badge.font.name = 'Calibri'
    run_badge.font.size = Pt(11.5)
    run_badge.font.color.rgb = RGBColor(100, 110, 120)
    run_badge.italic = True
    
    # Author details in a beautiful center alignment
    details_p = doc.add_paragraph()
    details_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_p.paragraph_format.space_before = Pt(40)
    details_p.paragraph_format.space_after = Pt(4)
    
    def add_meta_line(p, label, val):
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True
        r_lbl.font.size = Pt(10)
        r_lbl.font.color.rgb = RGBColor(10, 37, 64)
        r_val = p.add_run(f"{val}\n")
        r_val.font.size = Pt(10)
        r_val.font.color.rgb = RGBColor(51, 51, 51)
        
    add_meta_line(details_p, "Author / Developer", "Prathamesh Sawarkar (Registration ID: 12509401)")
    add_meta_line(details_p, "Institutional Scope", "Lovely Professional University (LPU) - CSE/IT Department")
    add_meta_line(details_p, "Contact", "prathameshsawarkar1@gmail.com")
    add_meta_line(details_p, "Document Version", "v2.0 (Industry-Grade Specification)")
    add_meta_line(details_p, "Target Platform", "FastAPI Backend + Next.js App Router Frontend + Supabase DB")
    
    doc.add_page_break()
    
    # Setup Header & Footer for rest of document
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header_p.add_run("CodeCanvas (LPU CodeViz) — Product Requirements Document (PRD)")
    run_h.font.name = 'Arial'
    run_h.font.size = Pt(8.5)
    run_h.font.color.rgb = RGBColor(120, 130, 140)
    
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run("Page ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9.0)
    run_f.font.color.rgb = RGBColor(120, 130, 140)
    add_page_number(run_f)
    
    # ─── DOCUMENT CONTROL METADATA ───
    add_custom_heading(doc, "Document Control & Metadata", level=1)
    
    table_meta = doc.add_table(rows=6, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    table_meta.columns[0].width = Inches(2.2)
    table_meta.columns[1].width = Inches(4.3)
    
    hdr_meta = table_meta.rows[0].cells
    hdr_meta[0].text = "Field"
    hdr_meta[1].text = "Details"
    
    meta_rows = [
        ("Product Name", "CodeCanvas (LPU CodeViz)"),
        ("Document Version", "v2.0 (Industry-Grade Specification)"),
        ("Status", "Approved"),
        ("Owner / Author", "Prathamesh Sawarkar (Registration ID: 12509401)"),
        ("Target Institutions", "Lovely Professional University (LPU) - CSE/IT Department")
    ]
    for idx, (fld, det) in enumerate(meta_rows, start=1):
        table_meta.rows[idx].cells[0].text = fld
        table_meta.rows[idx].cells[1].text = det
    style_table(table_meta)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    add_body_paragraph(doc, "Revision History:")
    table_rev = doc.add_table(rows=3, cols=4)
    table_rev.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_rev.autofit = False
    table_rev.columns[0].width = Inches(1.0)
    table_rev.columns[1].width = Inches(1.2)
    table_rev.columns[2].width = Inches(1.8)
    table_rev.columns[3].width = Inches(2.5)
    
    hdr_rev = table_rev.rows[0].cells
    hdr_rev[0].text = "Version"
    hdr_rev[1].text = "Date"
    hdr_rev[2].text = "Author"
    hdr_rev[3].text = "Description of Changes"
    
    revs = [
        ("v1.0", "July 7, 2026", "Prathamesh Sawarkar", "Initial draft & MVP features definition."),
        ("v2.0", "August 3, 2026", "Prathamesh Sawarkar", "Expanded to B2B enterprise specs, detailed API schemas, NAAC audits, DPDP Act 2023 controls, and Razorpay GST invoicing flows.")
    ]
    for idx, (ver, dt, auth, desc) in enumerate(revs, start=1):
        row = table_rev.rows[idx].cells
        row[0].text = ver
        row[1].text = dt
        row[2].text = auth
        row[3].text = desc
    style_table(table_rev)
    
    doc.add_page_break()
    
    # ─── SECTION 1: EXECUTIVE SUMMARY ───
    add_custom_heading(doc, "1. Executive Summary & Product Vision", level=1)
    
    add_body_paragraph(doc, 
        "CodeCanvas is an AI-powered code visualization and interactive tutoring SaaS platform designed specifically for higher education engineering colleges and universities in India. At its core, CodeCanvas bridges the critical gap between abstract programming concepts (memory models, variable state mutations, sorting algorithms, and database operations) and practical execution. It replaces standard command-line debuggers (like GDB/PDB) with an animated visual canvas, reducing student drop-out rates in computer science labs while automating accreditation audit documentation (NAAC/NBA).",
        bold_prefix="Product Vision: "
    )
    
    add_body_paragraph(doc, 
        "By replacing traditional, complex local debuggers (like GDB or PDB) with an intuitive visual canvas, and substituting high-risk code execution servers with a secure, serverless AI execution emulator (llama-3.3-70b-versatile via Groq), CodeCanvas offers a secure, highly engaging educational platform. Furthermore, the platform acts as an administrative asset for institutions. It monitors student progress, logs coding telemetry, and compiles compliance reports directly mapped to National Assessment and Accreditation Council (NAAC) and National Board of Accreditation (NBA) criteria."
    )
    
    add_callout_box(doc, 
        "CodeCanvas solves two primary problems simultaneously: (1) it decreases student failure rates in foundational programming courses (Data Structures and Algorithms) by making memory models visible, and (2) it automates hundreds of faculty hours spent compiling accreditation paperwork (NAAC Criterion 2.3) through automatic student progression export.",
        title="CORE VALUE PROPOSITION", color="0A2540"
    )
    
    # ─── SECTION 2: METRICS & KPIs ───
    add_custom_heading(doc, "2. Key Performance Indicators (KPIs) & Product Success Metrics", level=1)
    
    add_body_paragraph(doc, 
        "To ensure product efficacy, system performance, and institutional alignment, we trace key performance metrics across three domains:"
    )
    
    table_kpi = doc.add_table(rows=10, cols=3)
    table_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_kpi.autofit = False
    table_kpi.columns[0].width = Inches(1.8)
    table_kpi.columns[1].width = Inches(3.2)
    table_kpi.columns[2].width = Inches(1.5)
    
    hdr_kpi = table_kpi.rows[0].cells
    hdr_kpi[0].text = "Metric Domain"
    hdr_kpi[1].text = "Key Performance Indicator (KPI)"
    hdr_kpi[2].text = "Target Threshold"
    
    kpis = [
        ("Student Engagement", "Daily Active Users / Monthly Active Users (DAU/MAU)", ">= 40% cohort size"),
        ("Student Engagement", "Trace Completion Rate (steps completed vs started)", ">= 85%"),
        ("Student Engagement", "Retention & Coding Streak Maintenance (>= 5 days active)", ">= 30%"),
        ("Pedagogical Outcome", "Concept Comprehension Gain (Pre- vs. Post-Test difference)", ">= 25% improvement"),
        ("Pedagogical Outcome", "AI Grading Correlation Coefficient (r) with Instructor grades", "r >= 0.85"),
        ("Pedagogical Outcome", "AI Chat Session Self-Resolution Rate (no teacher escalation)", ">= 90%"),
        ("Technical Performance", "API Call Latency (llama-3.3-70b-versatile via Groq response)", "<= 2.0s (95th percentile)"),
        ("Technical Performance", "Visual Step Animation Transition Smoothness", ">= 60 fps"),
        ("Technical Performance", "SLA System Availability / Core Platform Uptime", "99.99%")
    ]
    for idx, (dom, ind, trg) in enumerate(kpis, start=1):
        row = table_kpi.rows[idx].cells
        row[0].text = dom
        row[1].text = ind
        row[2].text = trg
    style_table(table_kpi)
    
    doc.add_page_break()
    
    # ─── SECTION 3: PERSONAS ───
    add_custom_heading(doc, "3. Personas & Stakeholder User Scenarios", level=1)
    
    add_body_paragraph(doc, 
        "The system's user experience is tailored around three key roles within the university framework:"
    )
    
    add_body_paragraph(doc, 
        "1. Aarav (Engineering Student): Needs to visualize recursion stacks, variables, and pointers to pass exams. He works late at night and requires immediate help when stuck.",
        bold_prefix="Student Persona: ", list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "2. Dr. Meera (CS Department Lab Professor): Handles hundreds of students and multiple course files. She wants to set coding tasks, auto-grade traces, and export student performance records without manual paperwork.",
        bold_prefix="Faculty Persona: ", list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "3. Dean Dr. Sharma (Dean of Academics): Focuses on improving university results, maintaining security certifications, and securing high scores on accreditation evaluations (NAAC/NBA).",
        bold_prefix="Administrator Persona: ", list_style="List Bullet"
    )
    
    # ─── SECTION 4: MOSCOW PRIORITIZATION ───
    add_custom_heading(doc, "4. MoSCoW Feature Matrix & Prioritization", level=1)
    
    table_moscow = doc.add_table(rows=12, cols=4)
    table_moscow.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_moscow.autofit = False
    table_moscow.columns[0].width = Inches(1.8)
    table_moscow.columns[1].width = Inches(2.7)
    table_moscow.columns[2].width = Inches(1.1)
    table_moscow.columns[3].width = Inches(0.9)
    
    hdr_mos = table_moscow.rows[0].cells
    hdr_mos[0].text = "Feature Module"
    hdr_mos[1].text = "Functional Specifications & Core Capabilities"
    hdr_mos[2].text = "Target Persona"
    hdr_mos[3].text = "MoSCoW Category"
    
    feats = [
        ("Monaco Editor Integration", "Embeds a robust code editor with syntax highlighting, autocomplete, and auto-indentation for C, C++, Python, and SQL.", "Student", "Must-Have"),
        ("AI-Simulated Visualizer Engine", "FastAPI backend packages code and queries Groq's llama-3.3-70b-versatile model. Returns structured JSON containing step-by-step program states, preventing dangerous server-side execution of student scripts.", "Student", "Must-Have"),
        ("10 Animated Visualizer Canvases", "React-based custom visualizers animating states for: Arrays, Sorting, Stacks, Queues, Linked Lists, Binary Trees, Graphs, Recursion, SQL tables, and general Variable Boards.", "Student", "Must-Have"),
        ("Interactive Graph Creator", "Canvas layout enabling students to click to spawn vertices and drag connections to draw custom graphs. Automatically compiles Python graph adjacency lists.", "Student", "Must-Have"),
        ("Contextual AI Tutor Sidebar", "Slide-out AI Chatbot matching student query against their specific code and exact execution step, pointing out mistakes without giving direct answers.", "Student", "Must-Have"),
        ("One-Click Shareable Links", "Encodes code/visualizer state into a base64 string or unique DB slug, facilitating instant link sharing among students and professors.", "Student / Teacher", "Must-Have"),
        ("Multi-Tenant Classrooms", "Enables teachers to spawn unique classrooms with individual invite codes. Students input invite codes to enroll, populating a roster index.", "Teacher / Student", "Should-Have"),
        ("Assignments & Deadlines", "Enables professors to publish template code assignments with specific deadlines, maximum XP values, and course code links.", "Teacher / Student", "Should-Have"),
        ("Accreditation Export Engine", "One-click CSV/PDF generator that prints student enrollment rosters, cumulative grades, streak days, and syllabus completion mapped to NAAC criteria.", "HOD / Dean", "Should-Have"),
        ("LTI 1.3 LMS Passback", "Integration with standard LMS providers (Canvas, Blackboard, Moodle) to automatically sync classroom rosters and send back assignment grades.", "Teacher", "Could-Have"),
        ("Cryptographic Certificate Modal", "Generates NAAC-ready verified credentials and digital certifications for students showing topic completion, shareable on LinkedIn.", "Student", "Could-Have")
    ]
    for idx, (f_name, f_desc, f_user, f_scope) in enumerate(feats, start=1):
        row = table_moscow.rows[idx].cells
        row[0].text = f_name
        row[1].text = f_desc
        row[2].text = f_user
        row[3].text = f_scope
    style_table(table_moscow)
    
    doc.add_page_break()
    
    # ─── SECTION 5: TECHNICAL CONTRACTS ───
    add_custom_heading(doc, "5. Technical Stack & API Request-Response Contracts", level=1)
    
    add_body_paragraph(doc, 
        "CodeCanvas utilizes a Next.js 14 frontend, a FastAPI backend, a Supabase PostgreSQL database, and llama-3.3-70b-versatile via Groq. Below are the precise JSON integration contracts for core operations."
    )
    
    add_custom_heading(doc, "5.1 Visualizer Trace API (POST /api/trace)", level=2)
    add_body_paragraph(doc, "Request Payload:")
    add_code_block(doc, """{
  "code": "def search(arr, target):\\n    for i in range(len(arr)):\\n        if arr[i] == target:\\n            return i\\n    return -1\\nsearch([10, 20, 30], 20)",
  "lang": "python",
  "data_structure": "sorting"
}""")
    
    add_body_paragraph(doc, "Response Payload:")
    add_code_block(doc, """{
  "success": true,
  "data_structure_type": "sorting",
  "steps": [
    {
      "line_number": 1,
      "explanation": "Function 'search' is defined. Arguments are mapped: arr=[10, 20, 30], target=20.",
      "variables": {
        "arr": [10, 20, 30],
        "target": 20
      },
      "data_structure_state": {
        "type": "sorting",
        "array_state": [10, 20, 30],
        "highlighted_indices": [],
        "operation": "initialize"
      }
    },
    {
      "line_number": 4,
      "explanation": "Checking condition: arr[1] (20) == target (20). This evaluates to True. Returning index 1.",
      "variables": {
        "arr": [10, 20, 30],
        "target": 20,
        "i": 1
      },
      "data_structure_state": {
        "type": "sorting",
        "array_state": [10, 20, 30],
        "highlighted_indices": [1],
        "operation": "return"
      }
    }
  ]
}""")

    add_custom_heading(doc, "5.2 AI Tutor Chat API (POST /api/tutor/chat)", level=2)
    add_body_paragraph(doc, "Request Payload:")
    add_code_block(doc, """{
  "code": "def search(arr, target):\\n    for i in range(len(arr)):\\n        if arr[i] == target:\\n            return i\\n    return -1",
  "current_line": 3,
  "data_structure": "sorting",
  "variables": {
    "arr": [10, 20, 30],
    "target": 20,
    "i": 0
  },
  "user_message": "Why does the loop check arr[0] first?"
}""")
    
    add_body_paragraph(doc, "Response Payload:")
    add_code_block(doc, """{
  "response": "In Python, arrays and lists are 0-indexed. This means range(len(arr)) starts checking from index 0. Look at your variables sidebar: what is the current value of i?",
  "hint_type": "conceptual",
  "suggested_actions": ["check_variables_sidebar", "step_next"]
}""")
    
    doc.add_page_break()
    
    # ─── SECTION 6: DATABASE SCHEMA ───
    add_custom_heading(doc, "6. Relational Database Schema Specifications", level=1)
    add_body_paragraph(doc, 
        "The relational database schema is configured in PostgreSQL on Supabase. Row Level Security (RLS) is enabled on all tables, restricting students to their own data while allowing teachers to view enrollments and progress telemetry under their specific classes."
    )
    
    db_tables = [
        ("profiles", "Extends auth.users to store student/teacher profile data, XP, and active streaks.", 
         "id (UUID, PK), full_name (TEXT), avatar_url (TEXT), role (TEXT: student/teacher/hod/admin), school_id (TEXT, default 'cse'), xp (INT), streak_days (INT), last_active (DATE), created_at (TIMESTAMPTZ)"),
        ("classrooms", "Stores classroom details created by teachers, generating invite codes for student enrollment.", 
         "id (UUID, PK), teacher_id (UUID, FK -> profiles.id), name (TEXT), course_code (TEXT), school_id (TEXT), invite_code (TEXT, UNIQUE), description (TEXT), created_at (TIMESTAMPTZ)"),
        ("enrollments", "Links students to classrooms. Enforces unique student enrollment per class.", 
         "id (UUID, PK), classroom_id (UUID, FK -> classrooms.id), student_id (UUID, FK -> profiles.id), enrolled_at (TIMESTAMPTZ)"),
        ("assignments", "Stores assignments published by teachers for specific classrooms.", 
         "id (UUID, PK), classroom_id (UUID, FK -> classrooms.id), title (TEXT), description (TEXT), algorithm (TEXT), sample_code (TEXT), lang (TEXT), deadline (TIMESTAMPTZ), max_xp (INT), created_at (TIMESTAMPTZ)"),
        ("submissions", "Contains student code submissions, full execution trace states JSON, AI grades, and teacher grade overrides.", 
         "id (UUID, PK), assignment_id (UUID, FK -> assignments.id), student_id (UUID, FK -> profiles.id), code (TEXT), steps_json (JSONB), explanation (TEXT), ai_grade (INT), ai_feedback (TEXT), teacher_grade (INT), submitted_at (TIMESTAMPTZ)"),
        ("trace_history", "Logs telemetry of every code tracing event to construct student learning curves.", 
         "id (UUID, PK), user_id (UUID, FK -> profiles.id), lang (TEXT), code (TEXT), steps_json (JSONB), data_structure (TEXT), school_id (TEXT), step_count (INT), created_at (TIMESTAMPTZ)")
    ]
    
    table_db = doc.add_table(rows=7, cols=3)
    table_db.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_db.autofit = False
    table_db.columns[0].width = Inches(1.2)
    table_db.columns[1].width = Inches(2.2)
    table_db.columns[2].width = Inches(3.1)
    
    hdr_db = table_db.rows[0].cells
    hdr_db[0].text = "Table Name"
    hdr_db[1].text = "Role & Description"
    hdr_db[2].text = "Key Schema Attributes & Types"
    
    for idx, (t_name, t_desc, t_schema) in enumerate(db_tables, start=1):
        row = table_db.rows[idx].cells
        row[0].text = t_name
        row[1].text = t_desc
        row[2].text = t_schema
    style_table(table_db)
    
    doc.add_page_break()
    
    # ─── SECTION 7: ACCREDITATION & COMPLIANCE ───
    add_custom_heading(doc, "7. Educational Accreditation & Legal Compliance", level=1)
    
    add_custom_heading(doc, "7.1 NAAC / NBA Accreditation Alignment", level=2)
    add_body_paragraph(doc, 
        "Accreditation boards (NAAC and NBA) audit universities based on structured quality metrics. CodeCanvas serves as a data collector, automatically gathering student telemetry and outputting CSV/PDF audit documents:",
        bold_prefix="Accreditation Features: "
    )
    add_body_paragraph(doc, 
        "1. Experiential ICT Integration (NAAC Criterion 2.3.1): CodeCanvas satisfies this requirement by replacing traditional blackboard explanations with an interactive software visualizer that promotes active coding.",
        list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "2. Slow vs. Fast Learners Identification (NAAC Criterion 2.2.1): The analytics panel plots user progress, identifying students who complete assignments quickly (fast learners) versus students who take many tracing steps on basic loops (slow learners), enabling faculty to provide targeted support.",
        list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "3. CO-PO Attainment Analysis (NBA Criteria): Lab assessments can be linked directly to specific Course Outcomes (COs) and Program Outcomes (POs). Telemetry reports trace actual student grading data, proving outcomes achievement to auditors.",
        list_style="List Bullet"
    )
    
    add_custom_heading(doc, "7.2 Legal Compliance: Indian DPDP Act 2023", level=2)
    add_body_paragraph(doc, 
        "Under the Digital Personal Data Protection (DPDP) Act of 2023, universities must guarantee privacy protections for student data. CodeCanvas incorporates three core compliance mechanisms:",
        bold_prefix="Privacy Control: "
    )
    add_body_paragraph(doc, 
        "1. Notice & Explicit Consent (Section 5): On student onboarding, a mandatory consent window details what data is collected (profile name, email, lab submissions) and how it is used (class grading and institutional accreditation reports).",
        list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "2. Right to Correction and Erasure (Section 6): Students are provided with a self-service 'Erasure Portal' under their profile settings, allowing them to delete their profile and wipe all trace logs permanently.",
        list_style="List Bullet"
    )
    
    # ─── SECTION 8: COMMERCIAL MODEL & INVOICING ───
    add_custom_heading(doc, "8. Commercial Model, Billing & Razorpay Integration", level=1)
    
    add_custom_heading(doc, "8.1 Institutional Pricing & GST Structure", level=2)
    add_body_paragraph(doc, 
        "CodeCanvas operates under a SaaS licensing framework with billing tiers for individual subscriptions and institutional site licenses. All transactions within the Indian market include an 18% Goods and Services Tax (GST) (under SAC Code 997331: Licensing services for the right to use computer software).",
        bold_prefix="Pricing Model: "
    )
    
    table_price = doc.add_table(rows=4, cols=4)
    table_price.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_price.autofit = False
    table_price.columns[0].width = Inches(1.3)
    table_price.columns[1].width = Inches(1.5)
    table_price.columns[2].width = Inches(1.5)
    table_price.columns[3].width = Inches(2.2)
    
    hdr_pri = table_price.rows[0].cells
    hdr_pri[0].text = "License Tier"
    hdr_pri[1].text = "Base Price (INR)"
    hdr_pri[2].text = "GST (18%)"
    hdr_pri[3].text = "SaaS Capabilities Included"
    
    price_data = [
        ("Individual Student Monthly", "Rs. 254 / month", "Rs. 46 (Total: Rs. 300)", "Access to visualizer sandbox, AI Tutor chatbot, history logging, and share links."),
        ("Standard Departmental Annual", "Rs. 84,745 / year", "Rs. 15,255 (Total: Rs. 1,00,000)", "Up to 500 students, 10 teacher licenses, classroom boards, assignment grading, and progress dashboards."),
        ("Institutional Enterprise Unlimited", "Rs. 2,54,237 / year", "Rs. 45,763 (Total: Rs. 3,00,000)", "Unlimited students/faculty, SSO integrations, custom campus portal, custom database backups, and NAAC telemetry export modules.")
    ]
    for idx, (tier, base, gst, incl) in enumerate(price_data, start=1):
        row = table_price.rows[idx].cells
        row[0].text = tier
        row[1].text = base
        row[2].text = gst
        row[3].text = incl
    style_table(table_price)
    
    add_custom_heading(doc, "8.2 Payment Flow & Automated Billing", level=2)
    add_body_paragraph(doc, 
        "Subscriptions are processed via a Razorpay checkout integration, accepting UPI, Credit/Debit cards, and Netbanking. Upon successful webhook execution: (1) Supabase unlocks the user's role capabilities, (2) an automated PDF invoice is generated containing the customer's GSTIN, State Code (e.g., 03 for Punjab if billing LPU), and HSN/SAC code, and (3) the invoice is emailed to the payer, logging the record into the database."
    )
    
    doc.add_page_break()
    
    # ─── SECTION 9: ROADMAP ───
    add_custom_heading(doc, "9. MVP Implementation Roadmap & Phased Rollout", level=1)
    add_body_paragraph(doc, 
        "CodeCanvas development is structured across four release phases. The MVP core focuses on code visualization, while subsequent institutional rollouts introduce multi-tenancy and audit export capabilities."
    )
    
    table_roadmap = doc.add_table(rows=5, cols=4)
    table_roadmap.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_roadmap.autofit = False
    table_roadmap.columns[0].width = Inches(0.8)
    table_roadmap.columns[1].width = Inches(1.5)
    table_roadmap.columns[2].width = Inches(1.2)
    table_roadmap.columns[3].width = Inches(3.0)
    
    hdr_rod = table_roadmap.rows[0].cells
    hdr_rod[0].text = "Phase"
    hdr_rod[1].text = "Focus Area"
    hdr_rod[2].text = "Target Timeline"
    hdr_rod[3].text = "Deliverables & Key Milestones"
    
    roadmap_data = [
        ("Phase 1", "Core Engine & Visualizers", "Months 1 - 2 (Completed)", "Monaco Editor, Groq llama-3.3-70b simulation tracer, 10 animated canvases, variable boards, and live complexity profiling charts."),
        ("Phase 2", "Multi-Tenant Classroom", "Months 3 - 4 (Completed)", "Supabase DB schemas, student enrollment via invite codes, classroom rosters, assignment publishing, and student trace history tracking."),
        ("Phase 3", "AI Grading & Analytics", "Months 5 - 6 (Completed)", "AI-assisted grading feedback, teacher manual grade overrides, HOD dashboard panels, and NAAC telemetry export capabilities."),
        ("Phase 4", "Commerce & Compliance", "Months 7 - 8 (In Progress)", "Razorpay integration, automated 18% GST invoices, DPDP consent mechanisms, data deletion settings, and public SEO domain mapping.")
    ]
    for idx, (ph, foc, time, deliv) in enumerate(roadmap_data, start=1):
        row = table_roadmap.rows[idx].cells
        row[0].text = ph
        row[1].text = foc
        row[2].text = time
        row[3].text = deliv
    style_table(table_roadmap)
    
    # ─── SECTION 10: APPENDIX & GLOSSARY ───
    add_custom_heading(doc, "10. Appendix & Glossary", level=1)
    
    add_body_paragraph(doc, "DSA: Data Structures and Algorithms — the core theoretical computer science concepts taught in university labs.", bold_prefix="• ", space_after=4)
    add_body_paragraph(doc, "NAAC: National Assessment and Accreditation Council — an autonomous body that assesses and accredits higher education institutions in India.", bold_prefix="• ", space_after=4)
    add_body_paragraph(doc, "NBA: National Board of Accreditation — accredits engineering and technical education programs based on outcome achievements.", bold_prefix="• ", space_after=4)
    add_body_paragraph(doc, "DPDP Act 2023: Digital Personal Data Protection Act, 2023 — the comprehensive data privacy and protection law enacted by the Government of India.", bold_prefix="• ", space_after=4)
    add_body_paragraph(doc, "LLM: Large Language Model — deep learning algorithms (such as llama-3.3-70b) used to read, trace, and annotate programs.", bold_prefix="• ", space_after=4)
    add_body_paragraph(doc, "Execution Trace: The detailed record of line-by-line program state changes during execution.", bold_prefix="• ", space_after=4)
    add_body_paragraph(doc, "RLS: Row Level Security — database-level policies restricting data access based on authenticated user IDs.", bold_prefix="• ", space_after=12)
    
    # Save the document
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "prd.docx")
    doc.save(out_path)
    print(f"PRD Document generated successfully at: {out_path}")

if __name__ == "__main__":
    generate_prd()
