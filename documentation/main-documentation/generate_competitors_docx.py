import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color="0A2540", size="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color)
    tcBorders.append(left)
    for b in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'nil')
        tcBorders.append(node)
    tcPr.append(tcBorders)

def style_table(table, header_bg="0A2540", row_shd="F5F7FA"):
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_bg)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Arial'
                run.font.size = Pt(8.5) # Slightly smaller font to fit 9 columns
                
    for r_idx, row in enumerate(table.rows[1:], start=1):
        shd_color = row_shd if r_idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            if shd_color != "FFFFFF":
                set_cell_background(cell, shd_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:color'), 'E2E8F0')
            tcBorders.append(bottom)
            for b in ['left', 'right', 'top']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'nil')
                tcBorders.append(node)
            tcPr.append(tcBorders)

def add_custom_heading(doc, text, level, space_before=18, space_after=6):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    run = h.add_run(text)
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(10, 37, 64)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 128, 128)
        run.bold = True
    return h

def add_body_paragraph(doc, text, bold_prefix="", space_after=8, list_style=None):
    p = doc.add_paragraph(style=list_style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
        r_prefix.font.name = 'Calibri'
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = RGBColor(10, 37, 64)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_callout_box(doc, text, title="MARKET STRATEGY", color="008080"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4FBFB")
    set_cell_left_border(cell, color=color, size="36")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"{title}\n")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(0, 128, 128)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = RGBColor(51, 51, 51)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(12)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def generate_competitors():
    doc = Document()
    
    # Page Geometry Setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Normal style override
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # ─── COVER PAGE ───
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("COMPETITIVE ANALYSIS\n& DIFFERENTIATION MATRIX")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.color.rgb = RGBColor(10, 37, 64)
    run_title.bold = True
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(80)
    run_sub = sub_p.add_run("Strategic Profiling of 15 CS Education Competitors, Feature Maps & Moats")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0, 128, 128)
    run_sub.bold = True
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(20)
    run_meta = meta_p.add_run("Author: Prathamesh Sawarkar (Registration ID: 12509401)\n"
                              "Scope: CodeCanvas (LPU CodeViz)\n"
                              "Institutional Target: higher-Education B2B SaaS\n"
                              "Last Updated: August 11, 2026")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(10.5)
    run_meta.font.color.rgb = RGBColor(100, 110, 120)
    
    doc.add_page_break()
    
    # Setup Header & Footer for rest of document
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header_p.add_run("CodeCanvas — Competitor Analysis & Matrix")
    run_h.font.name = 'Arial'
    run_h.font.size = Pt(8.5)
    run_h.font.color.rgb = RGBColor(120, 130, 140)
    
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run("Page ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9.0)
    run_f.font.color.rgb = RGBColor(120, 130, 140)
    add_page_number(run_f)
    
    # ─── SECTION 1: PROFILES OF COMPETITORS ───
    add_custom_heading(doc, "1. Competitor Categorization & Profiles (15 Competitors)", level=1)
    add_body_paragraph(doc, 
        "We evaluate 15 competitors across four primary categories to outline market gaps and differentiate CodeCanvas:"
    )
    
    add_custom_heading(doc, "1.1 Visual Compilers & Tracing Tools", level=2)
    add_body_paragraph(doc, "Python Tutor: Great for single-file visual tracing, but compiles natively on servers (insecure), has no B2B features, and lacks interactive AI tutoring.", bold_prefix="• ")
    add_body_paragraph(doc, "Algorithm Visualizer: Open-source project that animates algorithms, but requires students to write complex visualization-specific library configurations manually.", bold_prefix="• ")
    add_body_paragraph(doc, "VisuAlgo: Renders algorithmic steps but is restricted to hardcoded preset examples; students cannot paste and compile their own scripts.", bold_prefix="• ")
    add_body_paragraph(doc, "USF Data Structure Visuals: Legacy academic tool that is Flash/JS based, with no code editor or custom script execution support.", bold_prefix="• ")
    
    add_custom_heading(doc, "1.2 Online Compilers & Cloud IDEs", level=2)
    add_body_paragraph(doc, "OnlineGDB: Renders standard text debug outputs (GDB logs) which are confusing for beginners. Lacks graphical animation panels and B2B classroom features.", bold_prefix="• ")
    add_body_paragraph(doc, "Replit: Highly collaborative VM-based shell, but is a generic developer tool with no specialized visualizers (sorting arrays, linked lists, SQL joins).", bold_prefix="• ")
    
    add_custom_heading(doc, "1.3 B2C EdTech & Career Accelerators", level=2)
    add_body_paragraph(doc, "Codecademy: Guided text paths, but visuals are static and hardcoded; cannot visual trace arbitrary student-written scripts.", bold_prefix="• ")
    add_body_paragraph(doc, "GeeksforGeeks: Large coding library but explanations are static text and diagrams, lacking dynamic visual debugger compilers.", bold_prefix="• ")
    add_body_paragraph(doc, "Coding Ninjas: Indian B2C courses; lacks a dynamic trace execution sandbox and deans' panels for quality audits.", bold_prefix="• ")
    add_body_paragraph(doc, "Scaler Academy: Mentor-led career prep; lacks automated visualizers for arbitrary code and B2B NAAC compliance exports.", bold_prefix="• ")
    add_body_paragraph(doc, "LeetCode: Practice and recruitment portal; code evaluation is pass/fail only, lacking step-by-step visual memory trace playback and AI tutor hints.", bold_prefix="• ")
    
    add_custom_heading(doc, "1.4 B2B LMS & Enterprise Assessment Engines", level=2)
    add_body_paragraph(doc, "Moodle LMS: General B2B university LMS; lacks built-in code compilers, visualizers, and coding-specific AI tutors.", bold_prefix="• ")
    add_body_paragraph(doc, "Google Classroom: Homework assignment management; general purpose, with no compiler sandbox, visual tracing, or accreditation analytics.", bold_prefix="• ")
    add_body_paragraph(doc, "HackerRank for Education: B2B assessment test engine focusing on evaluation outcomes; lacks live memory-trace visualizers and tutor sidebars.", bold_prefix="• ")
    add_body_paragraph(doc, "VS Code Codespaces / GitHub Education: Browser-based VS Code instance; interface is highly intimidating and complex for absolute beginners.", bold_prefix="• ")
    
    doc.add_page_break()
    
    # ─── SECTION 2: COMPARISON MATRIX ───
    add_custom_heading(doc, "2. Feature Comparison Matrix", level=1)
    
    table_matrix = doc.add_table(rows=13, cols=9)
    table_matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_matrix.autofit = False
    
    col_widths = [Inches(1.8), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.6)]
    for row in table_matrix.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    hdr = table_matrix.rows[0].cells
    hdr[0].text = "Feature / Capability"
    hdr[1].text = "CodeCanvas"
    hdr[2].text = "Python Tutor"
    hdr[3].text = "VisuAlgo"
    hdr[4].text = "OnlineGDB"
    hdr[5].text = "LeetCode"
    hdr[6].text = "Replit"
    hdr[7].text = "Moodle"
    hdr[8].text = "GFG"
    
    matrix_data = [
        ("Monaco Code Editor Shell", "Yes", "Yes", "No", "Yes", "Yes", "Yes", "No", "Yes"),
        ("Arbitrary Code Tracing", "Yes", "Yes", "No", "Yes", "No", "Yes", "No", "No"),
        ("10 Animated Canvases", "Yes", "No", "Yes", "No", "No", "No", "No", "No"),
        ("Empathetic AI Tutor", "Yes", "No", "No", "No", "No", "No", "No", "No"),
        ("Zero RCE Security", "Yes", "No", "Yes", "No", "No", "No", "Yes", "Yes"),
        ("B2B Invite Classrooms", "Yes", "No", "No", "No", "No", "Yes", "Yes", "No"),
        ("Automated AI Grader", "Yes", "No", "No", "No", "Yes", "No", "No", "Yes"),
        ("NAAC Criterion 2.3/2.2", "Yes", "No", "No", "No", "No", "No", "No", "No"),
        ("NBA CO-PO Maps", "Yes", "No", "No", "No", "No", "No", "No", "No"),
        ("Gamified Streaks & XP", "Yes", "No", "No", "No", "Yes", "Yes", "No", "Yes"),
        ("Vernacular UI selector", "Yes", "No", "No", "No", "No", "No", "Yes", "No"),
        ("LTI 1.3 LMS Grade Sync", "Yes", "No", "No", "No", "No", "Yes", "Yes", "No")
    ]
    for idx, row_data in enumerate(matrix_data, start=1):
        row = table_matrix.rows[idx].cells
        for col_idx, text in enumerate(row_data):
            row[col_idx].text = text
            # Format Yes as bold and blue-green for CodeCanvas, standard for others
            if col_idx == 1 and text == "Yes":
                for paragraph in row[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 128, 128)
    style_table(table_matrix)
    
    doc.add_page_break()
    
    # ─── SECTION 3: KEY DEFENSIVE MOATS ───
    add_custom_heading(doc, "3. Our Key Strategic Moats", level=1)
    
    add_body_paragraph(doc, 
        "CodeCanvas establishes distinct competitive barriers that prevent easy replication by generic compilers or edtech platforms:",
        bold_prefix="Defensive Barriers: "
    )
    
    add_body_paragraph(doc, 
        "1. Security Advantage (AI-Simulated Execution): Platforms like Python Tutor and Replit compile student code natively, creating high remote code execution (RCE) vulnerabilities and requiring complex sandboxing. CodeCanvas evaluates code scripts conceptually through Llama-3.3 on Groq, ensuring student scripts never run on server hardware.",
        list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "2. Administrative Advantage (Accreditation Integrations): Academic platforms (Moodle) and practice portals (LeetCode) cannot export metrics mapping to accreditation audits. CodeCanvas logs student trace telemetry, outputting certified documentation that satisfies NAAC Criterion 2.3 (experiential learning) and NBA outcomes tracking.",
        list_style="List Bullet"
    )
    add_body_paragraph(doc, 
        "3. Pedagogical Advantage (Empathetic Conversational Tutoring): Traditional compilers offer dry console errors. CodeCanvas provides an AI Tutor Sidebar that reads active line numbers and variable scopes, guiding students in regional languages to fix logical bugs without giving away answers.",
        list_style="List Bullet"
    )
    
    add_callout_box(doc, 
        "By merging security-compliant code execution, automated university accreditation audits, and context-aware multi-lingual AI tutoring, CodeCanvas occupies a unique, protected niche in B2B engineering higher education.",
        title="DIFFERENTIATION VALUE CONNOTATION", color="0A2540"
    )
    
    # Save the document
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "competitors_analysis.docx")
    doc.save(out_path)
    print(f"Competitor analysis document generated successfully at: {out_path}")

if __name__ == "__main__":
    generate_competitors()
